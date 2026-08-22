from __future__ import annotations

import hashlib
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from .baseline import (
    LOGICAL_PAGE_PATTERN,
    extract_logical_pages,
    extract_user_visible_text,
    inspect_docx_package,
)
from .hidden_markers import insert_hidden_markers_in_document
from .contracts import validate_deliverable_asset


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _unique_image_stream(path: Path, *, instance_id: str) -> BytesIO:
    """Return visually identical image bytes with a unique harmless trailer.

    python-docx deduplicates equal image blobs into one Word media part. Awesome
    1.0.1 uses that media-part identity as its reference id, so reusing a logo
    on multiple logical pages otherwise produces an output-path collision during
    project initialization. PNG/JPEG decoders ignore bytes after the image end;
    the trailer keeps pixels untouched while forcing a page-local media part.
    """
    trailer = f"\nPPT_WORD_INSTANCE:{instance_id}\n".encode("ascii")
    return BytesIO(path.read_bytes() + trailer)


def resolve_assignments(
    decisions: list[dict[str, object]], inventory: dict[str, object]
) -> dict[int, list[dict[str, object]]]:
    assets = {
        str(asset["asset_id"]): asset
        for asset in inventory.get("assets", [])
    }
    assignments: dict[int, list[dict[str, object]]] = {}
    for decision in decisions:
        selected = [
            str(value)
            for value in decision.get(
                "candidate_asset_ids",
                decision.get("selected_asset_ids", []),
            )
        ]
        if len(selected) > 16:
            raise ValueError(f"page {decision.get('page')} has more than 16 assets")
        if not selected:
            continue
        page = int(decision["page"])
        if decision.get("decision") != "ready":
            raise ValueError(f"selected assets require a ready decision: {page}")
        reviews = {
            str(review.get("asset_id")): review
            for review in decision.get("visual_reviews", [])
        }
        for asset_id in selected:
            review = reviews.get(asset_id)
            if (
                review is None
                or review.get("opened") is not True
                or review.get("visual_decision") != "accept"
                or not str(review.get("reason") or "").strip()
            ):
                raise ValueError(
                    f"asset requires an opened, accepted visual review: {page}: {asset_id}"
                )
        page_assets: list[dict[str, object]] = []
        for asset_id in selected:
            asset = assets.get(asset_id)
            if asset is None:
                raise ValueError(f"unknown asset: {asset_id}")
            if not asset.get("eligible"):
                raise ValueError(f"asset is not eligible: {asset_id}")
            contract = validate_deliverable_asset(asset)
            if not contract["valid"]:
                raise ValueError(
                    f"asset is not deliverable: {asset_id}: {contract['errors']}"
                )
            render_path = Path(str(asset.get("render_path") or ""))
            if not render_path.is_file():
                raise ValueError(f"asset render is missing: {asset_id}: {render_path}")
            page_assets.append(asset)
        assignments[page] = page_assets
    return assignments


def _director_role(asset: dict[str, object]) -> str:
    if asset.get("asset_role") == "identity_candidate":
        return "identity"
    asset_type = str(asset.get("asset_type") or "")
    if "chart" in asset_type or "table" in asset_type:
        return "factual_visual"
    if "diagram" in asset_type:
        return "structural_visual"
    return "scene_visual"


def _layout_rows(page_assets: list[dict[str, object]]) -> list[list[dict[str, object]]]:
    rows: list[list[dict[str, object]]] = []
    current: list[dict[str, object]] = []
    current_identity: bool | None = None
    for asset in page_assets:
        is_identity = _director_role(asset) == "identity"
        row_limit = 3 if current_identity else 2
        if current and (len(current) == row_limit or current_identity != is_identity):
            rows.append(current)
            current = []
            current_identity = None
        if not current:
            current_identity = is_identity
        current.append(asset)
    if current:
        rows.append(current)
    return rows


def _logical_labels(document: Document) -> list[tuple[int, object]]:
    labels: list[tuple[int, object]] = []
    for paragraph in document.paragraphs:
        match = LOGICAL_PAGE_PATTERN.fullmatch(paragraph.text.strip())
        if match is None or paragraph.style.style_id != "PaginationLabel":
            continue
        labels.append((int(match.group(1)), paragraph))
    return labels


def assemble_material_docx(
    source_docx: str | Path,
    decisions: list[dict[str, object]],
    inventory: dict[str, object],
    output_docx: str | Path,
    *,
    width_inches: float = 4.4,
    multi_width_inches: float = 3.1,
    identity_width_inches: float = 1.8,
) -> dict[str, object]:
    source = Path(source_docx)
    output = Path(output_docx)
    if source.resolve() == output.resolve():
        raise ValueError("output must differ from source")
    before_package = inspect_docx_package(source)
    before_logical = extract_logical_pages(source)
    assignments = resolve_assignments(decisions, inventory)
    document = Document(source)
    labels = _logical_labels(document)
    label_numbers = [page for page, _element in labels]
    if label_numbers != list(range(1, len(label_numbers) + 1)):
        raise ValueError(f"logical page labels are not consecutive: {label_numbers}")
    unknown_pages = sorted(set(assignments).difference(label_numbers))
    if unknown_pages:
        raise ValueError(f"logical page labels not found: {unknown_pages}")

    inserted: dict[int, list[dict[str, object]]] = {}
    for page, label_paragraph in labels:
        page_assets = assignments.get(page, [])
        if not page_assets:
            continue
        inserted[page] = []
        label_paragraph.paragraph_format.keep_with_next = True
        anchor = label_paragraph._p
        for row_number, row_assets in enumerate(_layout_rows(page_assets), start=1):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
            paragraph.paragraph_format.keep_together = True
            for column_number, asset in enumerate(row_assets, start=1):
                role = _director_role(asset)
                image_width = (
                    identity_width_inches
                    if role == "identity"
                    else multi_width_inches if len(row_assets) > 1 else width_inches
                )
                run = paragraph.add_run()
                render_path = Path(str(asset["render_path"]))
                image_stream = _unique_image_stream(
                    render_path,
                    instance_id=f"page-{page:03d}-row-{row_number:02d}-column-{column_number:02d}",
                )
                embedded_sha256 = hashlib.sha256(image_stream.getvalue()).hexdigest()
                run.add_picture(image_stream, width=Inches(image_width))
                inserted[page].append(
                    {
                        "asset_id": str(asset["asset_id"]),
                        "role": role,
                        "row": row_number,
                        "width_inches": image_width,
                        "sha256": _sha256(render_path),
                        "embedded_sha256": embedded_sha256,
                        "render_path": str(render_path.resolve()),
                    }
                )
            element = paragraph._p
            element.getparent().remove(element)
            anchor.addnext(element)
            anchor = element

    hidden_pages = insert_hidden_markers_in_document(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    if extract_user_visible_text(source) != extract_user_visible_text(output):
        raise ValueError("user-visible body text changed during assembly")
    after_package = inspect_docx_package(output)
    after_logical = extract_logical_pages(output)
    before_signatures = [page["signature"]["sha256"] for page in before_logical["pages"]]
    after_signatures = [page["signature"]["sha256"] for page in after_logical["pages"]]
    if before_signatures != after_signatures:
        raise ValueError("logical page text changed during assembly")
    if before_package["table_sha256"] != after_package["table_sha256"]:
        raise ValueError("table content changed during assembly")
    if before_package["comments_sha256"] != after_package["comments_sha256"]:
        raise ValueError("comment content changed during assembly")
    return {
        "source": str(source.resolve()),
        "output": str(output.resolve()),
        "source_sha256": _sha256(source),
        "output_sha256": _sha256(output),
        "logical_pages": len(labels),
        "inserted_pages": sorted(inserted),
        "inserted_assets": sum(len(items) for items in inserted.values()),
        "hidden_markers": len(hidden_pages),
        "visible_text_sha256": after_package["text_signature"]["sha256"],
        "logical_text_preserved": True,
        "table_count_preserved": True,
        "comment_count_preserved": True,
        "assignments": inserted,
    }
