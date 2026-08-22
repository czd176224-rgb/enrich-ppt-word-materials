from __future__ import annotations

import hashlib
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.oxml.ns import qn


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
LOGICAL_PAGE_PATTERN = re.compile(
    r"^第\s*(\d+)\s*页(?:\s*[·•|\-—]\s*.*)?$", re.IGNORECASE
)


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def make_text_signature(text: str, edge_chars: int = 80) -> dict[str, object]:
    normalized = re.sub(r"\s+", " ", text).strip()
    return {
        "normalized": normalized,
        "first": normalized[:edge_chars],
        "last": normalized[-edge_chars:] if normalized else "",
        "sha256": _sha256_bytes(normalized.encode("utf-8")),
        "characters": len(normalized),
    }


def extract_user_visible_text(path: str | Path) -> str:
    """Return body text while excluding nonprinting compatibility markers."""
    document = Document(path)
    blocks: list[str] = []
    for block in document.element.body.iterchildren():
        text_parts: list[str] = []
        for node in block.iter(qn("w:t")):
            hidden = False
            for run in node.iterancestors(qn("w:r")):
                properties = run.find(qn("w:rPr"))
                if properties is not None and properties.find(qn("w:vanish")) is not None:
                    hidden = True
                    break
            if not hidden and node.text:
                text_parts.append(node.text)
        text = "".join(text_parts).strip()
        if text:
            blocks.append(text)
    return "\n".join(blocks)


def _element_text(element: object) -> str:
    parts: list[str] = []
    for node in element.iter(qn("w:t")):
        hidden = False
        for run in node.iterancestors(qn("w:r")):
            properties = run.find(qn("w:rPr"))
            if properties is not None and properties.find(qn("w:vanish")) is not None:
                hidden = True
                break
        if not hidden and node.text:
            parts.append(node.text)
    return "".join(parts)


def _paragraph_style_id(element: object) -> str | None:
    style = element.find(
        f"./{{{WORD_NS}}}pPr/{{{WORD_NS}}}pStyle"
    )
    if style is None:
        return None
    return style.attrib.get(f"{{{WORD_NS}}}val")


def extract_logical_pages(
    path: str | Path, label_style_id: str = "PaginationLabel"
) -> dict[str, object]:
    document = Document(path)
    pages: list[dict[str, object]] = []
    current: dict[str, object] | None = None

    for block_index, element in enumerate(document.element.body.iterchildren()):
        text = _element_text(element).strip()
        tag = element.tag.rsplit("}", 1)[-1]
        is_label = (
            tag == "p"
            and _paragraph_style_id(element) == label_style_id
            and LOGICAL_PAGE_PATTERN.fullmatch(text) is not None
        )
        if is_label:
            if current is not None:
                current_text = "\n".join(current.pop("text_parts"))
                current["text"] = current_text
                current["signature"] = make_text_signature(current_text)
                pages.append(current)
            match = LOGICAL_PAGE_PATTERN.fullmatch(text)
            current = {
                "number": int(match.group(1)),
                "label": text,
                "start_block_index": block_index,
                "block_count": 1,
                "text_parts": [text],
            }
            continue
        if current is not None:
            current["block_count"] += 1
            if text:
                current["text_parts"].append(text)

    if current is not None:
        current_text = "\n".join(current.pop("text_parts"))
        current["text"] = current_text
        current["signature"] = make_text_signature(current_text)
        pages.append(current)

    numbers = [page["number"] for page in pages]
    return {
        "label_style_id": label_style_id,
        "logical_page_count": len(pages),
        "numbers": numbers,
        "is_sequential": numbers == list(range(1, len(numbers) + 1)),
        "pages": pages,
    }


def _count_comments(archive: zipfile.ZipFile) -> int:
    try:
        root = ElementTree.fromstring(archive.read("word/comments.xml"))
    except KeyError:
        return 0
    return len(root.findall(f"{{{WORD_NS}}}comment"))


def inspect_docx_package(path: str | Path) -> dict[str, object]:
    source = Path(path)
    raw = source.read_bytes()
    document = Document(source)

    with zipfile.ZipFile(source) as archive:
        document_xml = ElementTree.fromstring(archive.read("word/document.xml"))
        page_breaks = document_xml.findall(
            f".//{{{WORD_NS}}}br[@{{{WORD_NS}}}type='page']"
        )
        rendered_breaks = document_xml.findall(
            f".//{{{WORD_NS}}}lastRenderedPageBreak"
        )
        sections = document_xml.findall(f".//{{{WORD_NS}}}sectPr")
        media = [
            name
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        ]
        comments = _count_comments(archive)
        tables = document_xml.findall(f".//{{{WORD_NS}}}tbl")
        table_sha256 = [
            _sha256_bytes(ElementTree.tostring(table)) for table in tables
        ]
        try:
            comments_sha256 = _sha256_bytes(archive.read("word/comments.xml"))
        except KeyError:
            comments_sha256 = None

    visible_text = extract_user_visible_text(source)
    return {
        "path": str(source.resolve()),
        "bytes": len(raw),
        "sha256": _sha256_bytes(raw),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "table_sha256": table_sha256,
        "sections": len(sections),
        "manual_page_breaks": len(page_breaks),
        "rendered_page_breaks": len(rendered_breaks),
        "embedded_media": len(media),
        "comments": comments,
        "comments_sha256": comments_sha256,
        "text": visible_text,
        "text_signature": make_text_signature(visible_text),
    }


def capture_document_baseline(path: str | Path) -> dict[str, object]:
    from .word_com import capture_word_pages

    package = inspect_docx_package(path)
    logical_pages = extract_logical_pages(path)
    if not logical_pages["numbers"]:
        raise ValueError("logical page labels are missing")
    if not logical_pages["is_sequential"]:
        raise ValueError(
            f"logical page labels are not consecutive: {logical_pages['numbers']}"
        )
    physical_pages = capture_word_pages(path)
    return {
        "source_sha256": package["sha256"],
        "package": package,
        "logical_pages": logical_pages,
        "physical_pages": physical_pages,
    }
