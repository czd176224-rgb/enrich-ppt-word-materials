from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document
from docx.shared import Inches

from .baseline import LOGICAL_PAGE_PATTERN


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def insert_probe_images(
    source_docx: str | Path,
    assignments: dict[int, str | Path],
    output_docx: str | Path,
    *,
    width_inches: float = 0.15,
) -> dict[str, object]:
    source = Path(source_docx)
    output = Path(output_docx)
    document = Document(source)
    remaining = dict(assignments)
    inserted: list[int] = []

    for paragraph in document.paragraphs:
        match = LOGICAL_PAGE_PATTERN.fullmatch(paragraph.text.strip())
        if match is None or paragraph.style.style_id != "PaginationLabel":
            continue
        page_number = int(match.group(1))
        image_path = remaining.pop(page_number, None)
        if image_path is None:
            continue
        run = paragraph.add_run("\t")
        run.add_picture(str(Path(image_path).resolve()), width=Inches(width_inches))
        inserted.append(page_number)

    if remaining:
        raise ValueError(f"Logical page labels not found: {sorted(remaining)}")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return {
        "source": str(source.resolve()),
        "output": str(output.resolve()),
        "source_sha256": _sha256(source),
        "output_sha256": _sha256(output),
        "inserted_pages": inserted,
        "image_sha256": {
            str(page): _sha256(Path(assignments[page])) for page in inserted
        },
    }
