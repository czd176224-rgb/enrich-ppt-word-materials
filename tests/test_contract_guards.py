from pathlib import Path
from tempfile import TemporaryDirectory
from unittest import TestCase, mock

from docx import Document

from ppt_word_materials.assembly import assemble_material_docx, resolve_assignments
from ppt_word_materials.baseline import capture_document_baseline, inspect_docx_package
from ppt_word_materials.candidate_pool import build_candidate_pools
from ppt_word_materials.matching import build_shortlists


class ContractGuardTests(TestCase):
    def test_assembly_refuses_to_overwrite_source(self):
        source = Path("manuscript.docx")

        with self.assertRaisesRegex(ValueError, "output must differ from source"):
            assemble_material_docx(source, [], {"assets": []}, source)

    def test_selected_asset_requires_accepted_visual_review(self):
        decisions = [
            {
                "page": 1,
                "decision": "ready",
                "candidate_asset_ids": ["asset-1"],
                "visual_reviews": [],
            }
        ]

        with self.assertRaisesRegex(ValueError, "visual review"):
            resolve_assignments(decisions, {"assets": []})

    def test_selected_asset_requires_ready_decision(self):
        decisions = [
            {
                "page": 1,
                "decision": "ambiguous",
                "candidate_asset_ids": ["asset-1"],
                "visual_reviews": [
                    {
                        "asset_id": "asset-1",
                        "opened": True,
                        "visual_decision": "accept",
                        "reason": "Supports this page",
                    }
                ],
            }
        ]

        with self.assertRaisesRegex(ValueError, "ready"):
            resolve_assignments(decisions, {"assets": []})

    def test_baseline_rejects_missing_logical_page_labels(self):
        with TemporaryDirectory() as directory:
            source = Path(directory) / "manuscript.docx"
            document = Document()
            document.add_paragraph("No logical page labels")
            document.save(source)

            with mock.patch(
                "ppt_word_materials.word_com.capture_word_pages",
                return_value={"page_count": 1, "pages": []},
            ):
                with self.assertRaisesRegex(ValueError, "logical page labels"):
                    capture_document_baseline(source)

    def test_package_fingerprint_detects_table_content_change(self):
        with TemporaryDirectory() as directory:
            paths = [Path(directory) / f"table-{index}.docx" for index in (1, 2)]
            for path, value in zip(paths, ("before", "after")):
                document = Document()
                document.add_table(rows=1, cols=1).cell(0, 0).text = value
                document.save(path)

            before, after = map(inspect_docx_package, paths)
            self.assertEqual(before["tables"], after["tables"])
            self.assertNotEqual(before["table_sha256"], after["table_sha256"])

    def test_candidate_pool_removes_matching_delivery_visual_hashes(self):
        shortlist = {
            "page": 1,
            "page_need": {
                "candidate_range": [0, 8],
                "required_roles": [],
                "acceptable_roles": ["scene_visual"],
                "entities": [],
            },
            "candidates": [
                {
                    "asset_id": asset_id,
                    "director_role": "scene_visual",
                    "score": score,
                    "delivery_render_sha256": digest,
                    "delivery_visual_phash": "same-visual",
                }
                for asset_id, score, digest in (
                    ("asset-1", 1.0, "hash-1"),
                    ("asset-2", 0.9, "hash-2"),
                )
            ],
        }

        pools = build_candidate_pools([shortlist])
        self.assertEqual(pools[0]["candidate_asset_ids"], ["asset-1"])

    def test_joint_team_shortlist_allows_three_partner_logos(self):
        page = {
            "number": 1,
            "text": "第 1 页 · STORY LINE\n甲公司、乙公司、丙公司联合团队",
        }
        assets = [
            {
                "asset_id": entity,
                "asset_role": "identity_candidate",
                "identity_entities": [entity],
                "eligible": True,
                "asset_scope": "visual_object",
                "deliverable": True,
                "quality_status": "pass",
                "text": entity,
            }
            for entity in ("甲公司", "乙公司", "丙公司")
        ]

        candidates = build_shortlists([page], assets, top_k=16)[0]["candidates"]
        self.assertEqual(len(candidates), 3)


if __name__ == "__main__":
    import unittest

    unittest.main()
